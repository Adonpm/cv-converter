"""
CV Parser for extracting content from Whoz template CVs
"""
from docx import Document
import re
from pathlib import Path
from docx.oxml.ns import qn

class CVParser:
    def __init__(self):
        self.extracted_data = {}

    def parse_cv(self, file_path):
        """Parse a Whoz CV and extract structured data"""
        try:
            doc = Document(file_path)
            self.extracted_data = {
                "name": "",
                "current_position": "",
                "years_experience": "",
                "skills": [], # Needs to address
                "nationality": "",
                "date_of_birth": "",
                "education": [],
                "certifications": [],
                "professional_memberships": [],
                "software": [],
                "key_capabilities": "",
                "languages": [],
                "professional_experience": [],
                "professional_experience_summary": [],
                "external_experience": [],
                "contact": {}, # Needs to address
                "summary": "",
                "publications_or_awards": []
            }

            self._extract_content_from_tables(doc)
            return self.extracted_data

        except Exception as e:
            raise Exception(f"Error parsing CV: {str(e)}")

    def _extract_content_from_tables(self, doc):
        """Extract content from document tables"""
        for table_num, table in enumerate(doc.tables):
            self._process_table(table, table_num)

    def _process_table(self, table, table_num):
        """Process individual table based on its content"""
        for row_num, row in enumerate(table.rows):
            cells = []
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    cells.append(text)
                
            if not cells:
                continue

            # Process 1st table - Basic Information
            if table_num == 0:
                self._extract_basic_info(cells, row_num)

            # Process other tables based on section headers
            elif cells:
                potential_header = cells[0].upper()

                # Check if this is actually a header (no dates/numbers)
                if any(char.isdigit() for char in potential_header):
                    section_header = "PROFESSIONAL EXPERIENCE"
                else:
                    section_header = potential_header

                if "EDUCATION" in section_header:
                    if len(row.cells) >= 2:
                        right_cell = row.cells[1]
                        self._extract_education_cell(right_cell) # pass the actual right-hand _Cell
                        continue
                elif "CERTIFICATION" in section_header:
                    self._extract_certifications(cells)
                elif "PROFESSIONAL MEMBERSHIPS" in section_header or "ACCREDITATIONS" in section_header:
                    self._extract_memberships(cells)
                elif "SOFTWARE" in section_header:
                    self._extract_software(cells)
                elif "KEY CAPABILITIES" in section_header:
                    self._extract_key_capabilities(cells)
                elif "LANGUAGES" in section_header:
                    self._extract_languages(cells)
                elif "PROFESSIONAL EXPERIENCE" in section_header:
                    self._extract_professional_experience(cells, row_num, raw_row=row)
                elif "PUBLICATIONS" in section_header or "PROFESSIONAL AWARDS" in section_header:
                    self._extract_publications_professional_awards(cells)

    def _extract_basic_info(self, cells, row_num):
        """Extract basic information from first table"""
        if row_num == 0 and len(cells) >= 1:
            # Extract name from "LAST NAME First name: xxx yy"
            name_text = cells[0]
            if ":" in name_text:
                name = name_text.split(":")[-1].strip()

                if len(name.split()) >= 2:
                    last_name = name.split()[-1]
                    balance_name = " ".join(name.split()[:-1])
                    name = last_name + " " + balance_name
                self.extracted_data["name"] = name

            # Extract nationality if present
            if len(cells) >= 3 and "Nationality" in cells[2]:
                nationality_text = cells[2]
                if ":" in nationality_text:
                    nationality = nationality_text.split(":")[-1].strip()
                    if nationality != "(To be completed)":
                        self.extracted_data["nationality"] = nationality

        elif row_num == 1 and len(cells) >= 1:
            # Extract current postion
            position_text = cells[0]
            if "Current position:" in position_text:
                self.extracted_data["current_position"] = position_text.split(":")[-1].strip()

            # Extract date of birth if present
            if len(cells) >= 3 and "Date of birth" in cells[2]:
                dob_text = cells[2]
                if ":" in dob_text:
                    dob = dob_text.split(":")[-1].strip()
                    if not dob.startswith("(**/**/****)"):
                        self.extracted_data["date_of_birth"] = dob
            
        elif row_num == 2 and len(cells) >= 1:
            # Extract years of experience
            experience_text = cells[0]
            if "Years of experience:" in experience_text:
                self.extracted_data["years_experience"] =  experience_text.split(":")[-1].strip()
    
    def _extract_education_cell(self, cell):
        # Defensive: ensure we have a docx cell
        if not hasattr(cell, "tables") or not hasattr(cell, "paragraphs"):
            raise TypeError("Expected a python-docx cell object; got a non-cell value")
        
        entry_final = []
        for ti, tbl in enumerate(cell.tables):
            entry = {}
            for ri, row in enumerate(tbl.rows):
                # Expecting rows like ['School/University:', 'London School of Economics']
                if len(row.cells) >= 2:
                    label = row.cells[0].text.strip().rstrip(":")
                    value = " ".join(c.text.strip() for c in row.cells[1:] if c.text.strip())
                elif len(row.cells) == 1:
                    cell_text = row.cells[0].text.strip()
                    if ":" in cell_text:
                        label, value = [s.strip() for s in cell_text.split(":", 1)]
                    else:
                        label, value = cell_text, ""
                else:
                    continue

                if value and value != "(To be completed)":
                    entry[label] = value

            row_vals_dict = entry 
            entry_final.append(row_vals_dict)

        # Adding year extraction
        year_re = re.compile(r'^\s*(\d{4})(?:\s*-\s*(\d{4}))?\s*$')
        years_list = []
        for p in cell.paragraphs:
            txt = p.text.strip()
            m = year_re.match(txt)
            if not m:
                continue
            if m.group(2):
                years_list.append(f"{m.group(1)}-{m.group(2)}") # prints "2017-2020"
            else: 
                years_list.append(f"{m.group(1)}") # prints "2022"
        for i in range(len(years_list)):
            entry_final[i]["year"] = years_list[i]

        self.extracted_data["education"] = entry_final
  
    def _extract_certifications(self, cells):
        """Extract certification information"""
        if len(cells) >= 2:
            cert_text = cells[1].strip()
            if cert_text and cert_text != "(To be completed)":
                # Parse certification details
                cert_info = {}
                lines = [ln.strip() for ln in cert_text.split('\n') if ln.strip]
                for line in lines:
                    line = line.strip()
                    if "Obtention date:" in line:
                        cert_info["date"] = line.replace("Obtention date:", "").strip()
                    elif "Delivering entity:" in line:
                        cert_info["entity"] = line.replace("Delivering entity:", "").strip()
                    elif "Programme:" in line:
                        cert_info["programme"] = line.replace("Programme:", "").strip()
                    
                    if cert_info.get("date") and cert_info.get("entity") and cert_info.get("programme"):
                        self.extracted_data["certifications"].append(cert_info)
                        cert_info = {}
                
    def _extract_memberships(self, cells):
        """Extract professional memberships"""
        if len(cells) >= 2:
            membership_text = cells[1].strip()
            if membership_text and membership_text != "(To be completed)":
                # Parse membership details
                lines = [ln.strip() for ln in membership_text.split('\n') if ln.strip]
                membership_info = {} # dict for the current membership
                for line in lines:
                    if "Obtention date:" in line:
                        membership_info["date"] = line.replace("Obtention date:", "").strip()
                    elif "Programme:" in line:
                        membership_info["programme"] = line.replace("Programme:", "").strip()

                    if membership_info.get("programme") and membership_info.get("date"):
                        self.extracted_data["professional_memberships"].append(membership_info)
                        membership_info = {}

    def _extract_software(self, cells):
        """Extract software information"""
        if len(cells) >= 2:
            software_text = cells[1].strip()
            if software_text and software_text != "(To be completed)":
                self.extracted_data["software"].append(software_text)

    def _extract_key_capabilities(self, cells):
        """Extract key capabilities"""
        if len(cells) >= 2:
            capabilities_text = cells[1].strip()
            if capabilities_text and capabilities_text != "(To be completed)":
                self.extracted_data["key_capabilities"] = capabilities_text
                
                # Also extract summary from presentation section
                if "Presentation:" in capabilities_text:
                    presentation_part = capabilities_text.split("Presentation:")[-1].strip()
                    self.extracted_data["summary"] = presentation_part

    def _extract_languages(self, cells):
        """Extract language information"""
        if len(cells) >= 2:
            language_text = cells[1].strip()
            if language_text:
                # Parse language entries (e.g., "English: Native – C2")
                languages = language_text.split(',')
                for lang in languages:
                    lang = lang.strip()
                    if ':' in lang:
                        lang_name = lang.split(':')[0].strip()
                        lang_level = lang.split(':')[1].strip()
                        self.extracted_data["languages"].append({
                            "language": lang_name,
                            "level": lang_level
                        })

    @staticmethod
    def _cell_text_with_bullets(cell):
        """Extract cell text, prefixing Word-bulleted paragraphs with '• '"""
        lines = []
        for para in cell.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            pPr = para._element.find(qn('w:pPr'))
            is_bullet = pPr is not None and pPr.find(qn('w:numPr')) is not None
            lines.append(f" • {text}" if is_bullet else text)
        return "\n".join(lines)

    def _extract_professional_experience(self, cells, row_num, raw_row=None):
        """Extract professional experience"""
        # Skip header row
        if row_num == 0:
            return

        if len(cells) >= 2:
            # First cell contains duration and dates
            duration_text = cells[0].strip()
            # Use bullet-aware extraction if raw row cells are available
            if raw_row is not None:
                job_text = self._cell_text_with_bullets(raw_row.cells[1])
            else:
                job_text = cells[1].strip()

            if duration_text and job_text and job_text != "(To be completed)":
                experience_entry = {}
                experience_summary_entry = {}
                
                # Parse job details for company, role, location, etc.
                self._parse_job_details(experience_entry, experience_summary_entry, job_text, duration_text)
                
                self.extracted_data["professional_experience"].append(experience_entry)
                self.extracted_data["professional_experience_summary"].append(experience_summary_entry)

    def _compute_end_year(self, start_date, duration_period):
        """Compute the end year from a start date (expects a MM/YYYY substring)
        and a duration string like '10 months', '1 year 10 months', '2 years'.
        Correctly distinguishes years vs months instead of assuming the first
        number in the duration is always a year count."""
        year_match = re.findall(r'\b(?:19|20)\d{2}\b', start_date or "")
        if not year_match:
            return None
        start_year = int(year_match[0])

        month_of_start_match = re.search(r'(\d{1,2})/(\d{4})', start_date)
        start_month = int(month_of_start_match.group(1)) if month_of_start_match else 1

        y_match = re.search(r'(\d+)\s*year', duration_period or "")
        m_match = re.search(r'(\d+)\s*month', duration_period or "")
        years = int(y_match.group(1)) if y_match else 0
        months = int(m_match.group(1)) if m_match else 0

        if not y_match and not m_match:
            # No recognizable unit — fall back to treating the first number as years
            nums = re.findall(r'\d+', duration_period or "")
            years = int(nums[0]) if nums else 0

        total_month_index = start_year * 12 + (start_month - 1) + years * 12 + months
        return str(total_month_index // 12)

    def _parse_job_details(self, experience_entry, experience_summary_entry, job_text, duration_text):
        """Parse detailed job information"""
        if "Project title" in job_text:
            # Parse duration for start/end dates
            if "From" in duration_text and "Duration" in duration_text:
                parts = re.split(r'From|Duration:', duration_text)
                if len(parts) >= 2:
                    experience_entry["start_date"] = parts[1].strip()
                if len(parts) >= 3:
                    experience_entry["duration_period"] = parts[2].strip()
                if experience_entry["start_date"] is not None and experience_entry["duration_period"] is not None:
                    start_year = re.findall(r'\b(?:19|20)\d{2}\b', experience_entry["start_date"])[0]
                    end_year = self._compute_end_year(experience_entry["start_date"], experience_entry["duration_period"])
                    experience_entry["duration_period_year"] = f"{start_year} to {end_year}"

            elif "Duration" in duration_text:
                parts = re.split(r'Duration:', duration_text)
                if len(parts) >= 2:
                    experience_entry["start_date"] = parts[0].strip()
                    experience_entry["duration_period"] = parts[1].strip()
                if experience_entry["start_date"] is not None and experience_entry["duration_period"] is not None:
                    start_year = re.findall(r'\b(?:19|20)\d{2}\b', experience_entry["start_date"])[0]
                    end_year = self._compute_end_year(experience_entry["start_date"], experience_entry["duration_period"])
                    experience_entry["duration_period_year"] = f"{start_year} - {end_year}"

            if "Location" in job_text and "Client name" in job_text and "Detailed project description" in job_text and "Detailed description of the expert mission (on this project)" in job_text and "Technical expertise" in job_text:
                raw_parts = re.split(r'Project title:|Location:|Client name:|Detailed project description:|Detailed description of the expert mission \(on this project\):|Technical expertise:', job_text)
                parts = [part for part in raw_parts if part is not None]
                experience_entry["experience_header"] = parts[0].strip()
                experience_entry["project_title"] = parts[1].strip()
                experience_entry["location"] = parts[2].strip()
                experience_entry["client"] = parts[3].strip()
                experience_entry["project_description"] = parts[4].strip()
                experience_entry["expert_mission_description"] = parts[5].strip()
                experience_entry["technical_expertise"] = parts[6].strip()
        else:
            # Parse duration for start/end dates
            if "From" in duration_text and "Duration" in duration_text:
                parts = re.split(r'From|Duration:', duration_text)
                if len(parts) >= 2:
                    experience_summary_entry["start_date"] = parts[1].strip()
                if len(parts) >= 3:
                    experience_summary_entry["duration_period"] = parts[2].strip()
            elif "Duration" in duration_text:
                parts = re.split(r'Duration:', duration_text)
                if len(parts) >= 2:
                    experience_summary_entry["start_date"] = parts[0].strip()
                    experience_summary_entry["duration_period"] = parts[1].strip()

            if "Location" in job_text and "Description" in job_text:
                parts = re.split(r'Location:|Description:', job_text)
                header_lines = [ln.strip() for ln in parts[0].split('\n') if ln.strip()]
                experience_summary_entry["employee_company"] = header_lines[0] if header_lines else ""
                experience_summary_entry["employee_position"] = header_lines[1] if len(header_lines) >= 2 else ""
                experience_summary_entry["employee_location"] = parts[1].strip() if len(parts) > 1 else ""
                experience_summary_entry["employee_description"] = parts[2].strip() if len(parts) > 2 else ""

                # NEW: also capture this as an "external" (non-SYSTRA) project entry
                # so it can be rendered in {{PROJECTS}} with a red header.
                company = experience_summary_entry["employee_company"]
                if company and "systra" not in company.lower():
                    position = experience_summary_entry["employee_position"]
                    description = experience_summary_entry["employee_description"]
                    start_date = experience_summary_entry.get("start_date", "") or ""

                    years = re.findall(r'\b(?:19|20)\d{2}\b', start_date)
                    if len(years) >= 2:
                        date_range = f"{years[0]} to {years[1]}"
                    elif len(years) == 1:
                        date_range = years[0]
                    else:
                        date_range = start_date.strip()

                    header_parts = [company]
                    if position:
                        header_parts.append(position)
                    header_line = " - ".join(header_parts)
                    if date_range:
                        header_line += f", {date_range}"

                    self.extracted_data["external_experience"].append({
                        "header": header_line,
                        "description": description
                    })

    def _extract_publications_professional_awards(self, cells):
        """Extract Publications or Awards"""
        if len(cells) >= 2:
            publications_professional_awards_text = cells[1].strip()
            if publications_professional_awards_text and publications_professional_awards_text != "(To be completed)":
                self.extracted_data["publications_or_awards"].append(publications_professional_awards_text)
