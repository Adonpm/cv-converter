"""
Template manager for handling CV templates
"""
from docx import Document
from pathlib import Path
from utils.config import Config
import shutil
from docx.shared import Pt, RGBColor
import re
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement

class TemplateManager:
    def __init__(self):
        self.templates = {}
        self.load_templates()

    def load_templates(self):
        """Load available templates"""
        self.templates = {}
        templates_dir = Config.TEMPLATES_DIR

        if templates_dir.exists():
            for template_file in templates_dir.glob("*.docx"):
                template_name = template_file.stem
                self.templates[template_name] = str(template_file)

    def get_template_list(self):
        """Get list of available templates"""
        return list(self.templates.keys())
    
    def get_template_path(self, template_name):
        """Get the file path of a specific template"""
        return self.templates.get(template_name)
    
    def apply_template(self, template_name, cv_data, output_path, formatting_options=None):
        """Apply a template to the provided CV data"""
        template_path = self.get_template_path(template_name)
        if not template_path:
            raise ValueError(f"Template '{template_name}' not found.")
        
        try:
            # Load template
            doc = Document(template_path)

            # Apply CV data to the template
            self._populate_template(doc, cv_data, formatting_options or {})

            # Save output
            doc.save(output_path)
            return True
        
        except Exception as e:
            raise Exception(f"Error applying template: {str(e)}")

    def _populate_template(self, doc, cv_data, formatting_options):
        """Populate template with CV data"""

        # Handle {{PROJECTS}} first — replaces single paragraph with multiple
        # properly styled paragraphs instead of cramming all content into one
        for paragraph in doc.paragraphs:
            if '{{PROJECTS}}' in paragraph.text:
                self._insert_projects_as_paragraphs(paragraph, cv_data, formatting_options)
                break  # paragraph list is now modified, stop iterating

        # Replace placeholders in remaining paragraphs
        for paragraph in doc.paragraphs:
            self._replace_placeholders(paragraph, cv_data, formatting_options)
        
        # Replace placeholders inside text boxes (floating shapes)
        WPS_TXBX = '{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx'
        W_P = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'
        for shape in doc.element.body.iter(WPS_TXBX):
            for p_elem in shape.iter(W_P):
                # Use doc.element as parent to avoid NoneType 'part' error
                paragraph = Paragraph(p_elem, doc)
                self._replace_placeholders(paragraph, cv_data, formatting_options)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_placeholders(paragraph, cv_data, formatting_options)

        # Update headers and footers
        for section in doc.sections:
            self._update_header_footer(section, formatting_options)

    def _make_paragraph_element(self, style_val):
        """Create a new <w:p> element with the given Word style."""
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        pStyle.set(qn('w:val'), style_val)
        pPr.append(pStyle)
        p.append(pPr)
        return p

    def _make_run_element(self, text, bold=False, color=None, formatting_options=None, template_font_info=None):
        """Create a <w:r> run element with user-specified font or template font fallback.

        color: optional hex string (e.g. "C00000") to set explicit font color,
        used to flag external (non-SYSTRA) experience headers in red.
        """
        
        formatting_options = formatting_options or {}
        template_font_info = template_font_info or {}

        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        if color:
            c = OxmlElement('w:color')
            c.set(qn('w:val'), color)
            rPr.append(c)

        # Font family — user choice wins, else template font, else nothing
        font_name = formatting_options.get("font_family", "")
        if font_name:
            fonts = OxmlElement('w:rFonts')
            fonts.set(qn('w:cs'), font_name)
            fonts.set(qn('w:ascii'), font_name)
            fonts.set(qn('w:hAnsi'), font_name)
            rPr.append(fonts)
        elif template_font_info.get('name'):
            fonts = OxmlElement('w:rFonts')
            fonts.set(qn('w:cs'), template_font_info['name'])
            fonts.set(qn('w:ascii'), template_font_info['name'])
            fonts.set(qn('w:hAnsi'), template_font_info['name'])
            rPr.append(fonts)

        if bold:
            b = OxmlElement('w:b')
            rPr.append(b)

        # Font size — user choice wins, else template size, else nothing
        font_size = formatting_options.get("font_size", "")
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(font_size) * 2))
            rPr.append(sz)
        elif template_font_info.get('size'):
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(template_font_info['size'] / 12700) * 2))
            rPr.append(sz)

        r.append(rPr)
        
        t = OxmlElement('w:t')
        t.text = text
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        r.append(t)
        return r

    def _insert_projects_as_paragraphs(self, projects_paragraph, cv_data, formatting_options):
        """
        Replace the single {{PROJECTS}} paragraph with individual properly-styled
        paragraphs — one CV h3 bold header + one cv text description per project.
        This fixes the column overflow issue caused by dumping all projects into
        one paragraph.
        """
        # Read font info from {{PROJECTS}} placeholder before replacing
        template_font_info = {}
        if projects_paragraph.runs:
            first_run = projects_paragraph.runs[0]
            template_font_info = {
                'name': first_run.font.name,
                'size': first_run.font.size,
            }
        
        # If not on run, try paragraph style chain
        if not template_font_info.get('name'):
            style = projects_paragraph.style
            while style:
                if style.font.name:
                    template_font_info['name'] = style.font.name
                    break
                style = style.base_style

        if not template_font_info.get('size'):
            style = projects_paragraph.style
            while style:
                if style.font.size:
                    template_font_info['size'] = style.font.size
                    break
                style = style.base_style

        experience_list = cv_data.get("professional_experience", [])
        parent = projects_paragraph._element.getparent()
        projects_p = projects_paragraph._element
        insert_position = list(parent).index(projects_p)

        new_elements = []
        for exp in experience_list:
            header_parts = []
            if exp.get("experience_header"):
                header_parts.append(exp["experience_header"])
            if exp.get("client"):
                header_parts.append(exp["client"])
            if exp.get("duration_period_year"):
                header_parts.append(exp["duration_period_year"])
            header_text = ", ".join(header_parts)
            project_desc = exp.get("project_description", "").strip()
            mission_desc = exp.get("expert_mission_description", "").strip()

            # Combine with a blank line between them if both exist
            if project_desc and mission_desc:
                description = project_desc + "\n" + mission_desc
            elif project_desc:
                description = project_desc
            else:
                description = mission_desc

            if header_text:
                h_p = self._make_paragraph_element("CVh3")
                h_p.append(self._make_run_element(
                    header_text,
                    bold=True,
                    formatting_options=formatting_options,
                    template_font_info=template_font_info
                ))
                new_elements.append(h_p)
            if description:
                desc_lines = description.split('\n')
                for line in desc_lines:
                    line = line.strip()
                    if not line:
                        continue
                    is_bullet = line.startswith(' • ')
                    line_text = line[3:] if is_bullet else line

                    if is_bullet:
                        d_p = self._make_paragraph_element("cvtext")
                        d_p.append(self._make_run_element(
                            f" • {line_text}",   # <-- literal bullet char
                            bold=False,
                            formatting_options=formatting_options,
                            template_font_info=template_font_info
                        ))
                    else:
                        d_p = self._make_paragraph_element("cvtext")
                        d_p.append(self._make_run_element(
                            line_text,
                            bold=False,
                            formatting_options=formatting_options,
                            template_font_info=template_font_info
                        ))

                    new_elements.append(d_p)

        # NEW: append external (non-SYSTRA) experience after the SYSTRA projects,
        # same paragraph styles, but header rendered in red instead of bold.
        for ext in cv_data.get("external_experience", []):
            header_text = ext.get("header", "")
            description = ext.get("description", "")

            if header_text:
                h_p = self._make_paragraph_element("CVh3")
                h_p.append(self._make_run_element(
                    header_text,
                    bold=False,
                    color="C00000",
                    formatting_options=formatting_options,
                    template_font_info=template_font_info
                ))
                new_elements.append(h_p)
            if description:
                for line in description.split('\n'):
                    line = line.strip()
                    if not line:
                        continue
                    is_bullet = line.startswith(' • ')
                    line_text = line[3:] if is_bullet else line

                    d_p = self._make_paragraph_element("cvtext")
                    d_p.append(self._make_run_element(
                        f" • {line_text}" if is_bullet else line_text,
                        bold=False,
                        formatting_options=formatting_options,
                        template_font_info=template_font_info
                    ))
                    new_elements.append(d_p)

        for i, elem in enumerate(new_elements):
            parent.insert(insert_position + i, elem)
        parent.remove(projects_p)

        # Move the text box anchor paragraph to just after Para[0] ({{NAME}})
        # so it anchors to page 1 instead of the last page
        self._move_textbox_to_page_one(parent)

    def _move_textbox_to_page_one(self, body):
        """
        Find the paragraph containing the floating text box and move it
        to just after the first paragraph ({{NAME}}) so it anchors to page 1.
        """
        WPS_TXBX = '{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}txbx'
        
        # Find the paragraph containing the text box
        txbx_para = None
        for child in body:
            if child.tag == '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p':
                if any(el.tag == WPS_TXBX for el in child.iter()):
                    txbx_para = child
                    break
        
        if txbx_para is None:
            return  # No text box found, nothing to do
        
        # Remove from current position
        body.remove(txbx_para)
        
        # Insert at position 1 — just after the first paragraph ({{NAME}})
        body.insert(1, txbx_para)

    def _replace_placeholders(self, paragraph, cv_data, formatting_options):
        """Replace placeholders in paragraph text"""
        text = paragraph.text

        # If user selected blank for both, skip font override entirely
        apply_custom_font = bool(formatting_options.get("font_family") or formatting_options.get("font_size"))

        # Define placeholder mappings
        placeholders = {
            "{{NAME}}": (cv_data.get("name", "") or "").upper(),
            "{{TITLE}}": cv_data.get("current_position", "") or "",
            "{{SUMMARY}}": cv_data.get("summary", "") or "",
            "{{HOME}}": cv_data.get("nationality", "") or "No data available",
            "{{YEARS OF EXPERIENCE}}": cv_data.get("years_experience", "") or "", 
            "{{EXPERIENCE SUMMARY}}": self._format_professional_experience_summary(cv_data.get("professional_experience_summary", [])),
            "{{EDUCATION}}": self._format_education(cv_data.get("education", [])),
            "{{SKILLS}}": self._format_skills(cv_data.get("professional_experience", [])), 
            "{{MEMBERSHIPS}}": self._format_memberships(cv_data.get("professional_memberships", [])),
            "{{CERTIFICATIONS}}": self._format_certifications(cv_data.get("certifications", [])),
            "{{PROJECTS}}": self._format_professional_experience(cv_data.get("professional_experience", []), cv_data.get("external_experience", [])),
            "{{CLIENT}}": formatting_options.get("client_name", "") or "",
            "{{OPPORTUNITY}}": formatting_options.get("opportunity_name", "") or ""
        }

        # Replace placeholders
        for placeholder, value in placeholders.items():
            if placeholder not in text:
                continue

            if placeholder in text and placeholder != "{{PROJECTS}}":
                # Ensure value is a string, not None
                safe_value = str(value) if value is not None else ""

                if not safe_value.strip():
                    self._delete_paragraph(paragraph)
                    continue

                # Store formatting before clearing
                original_style = paragraph.style

                # Replace text directly in XML text nodes — never touches formatting
                W_T = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'
                full_text = "".join(
                    el.text or "" 
                    for el in paragraph._element.iter(W_T)
                )
                new_text = full_text.replace(placeholder, safe_value)

                # Write new text — split on \n and add line breaks
                W_BR = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br'
                t_elements = list(paragraph._element.iter(W_T))
                if t_elements:
                    first_t = t_elements[0]
                    parent_r = first_t.getparent()  # <w:r> element
                    
                    # Clear all existing <w:t> elements
                    for t in t_elements:
                        t.text = ""
                    
                    # Split new text by newlines and insert <w:br/> between them
                    lines = new_text.split('\n')
                    first_t.text = lines[0]
                    
                    for line in lines[1:]:
                        # Add line break
                        br = OxmlElement('w:br')
                        parent_r.append(br)
                        # Add text after break
                        t = OxmlElement('w:t')
                        t.text = line
                        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                        parent_r.append(t)

                # Restore formatting
                paragraph.style = original_style

                # Apply font formatting to all placeholders
                if placeholder != "{{NAME}}" and apply_custom_font:
                    if formatting_options.get("font_family"):
                        for r in paragraph.runs:
                            r.font.name = formatting_options["font_family"]
                            # Remove theme font attributes that override explicit font name
                            rFonts = r._element.find(qn('w:rPr') + '/' + qn('w:rFonts'))
                            if rFonts is None:
                                rPr = r._element.find(qn('w:rPr'))
                                if rPr is not None:
                                    rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is not None:
                                rFonts.attrib.pop(qn('w:asciiTheme'), None)
                                rFonts.attrib.pop(qn('w:hAnsiTheme'), None)
                                rFonts.attrib.pop(qn('w:cstheme'), None)
                    if formatting_options.get("font_size"):
                        for r in paragraph.runs:
                            r.font.size = Pt(int(formatting_options["font_size"]))

            elif placeholder in text and placeholder == "{{PROJECTS}}":
                projects_text = self._format_professional_experience(cv_data.get("professional_experience", []), cv_data.get("external_experience", []))

                if not projects_text:
                    self._delete_paragraph(paragraph)
                    continue
                # Parse and add formatted text
                self._add_formatted_text(paragraph, projects_text, formatting_options)

    def _delete_paragraph(self, paragraph):
        '''
        Remove the paragraph from the document entirely
        '''
        p = paragraph._element
        parent = p.getparent()
        if parent is not None:
            parent.remove(p)

    def _add_formatted_text(self, paragraph, text, formatting_options=None):
        """Add text with formatting markers to paragraph"""
        
        formatting_options = formatting_options or {}

        # If user selected blank for both, skip font override entirely
        apply_custom_font = bool(formatting_options.get("font_family") or formatting_options.get("font_size"))

        original_style = paragraph.style
        original_font_info = {}
    
        # Get font info from the first run (if exists)
        if paragraph.runs:
            first_run = paragraph.runs[0]
            original_font_info = {
                'name': first_run.font.name,
                'size': first_run.font.size,
                'color': first_run.font.color.rgb if first_run.font.color and first_run.font.color.rgb else None,
                'italic': first_run.italic
            }
        
        # Set paragraph style BEFORE adding runs
        paragraph.style = original_style

        # Clear existing text and formatting
        paragraph.clear()

        def apply_common_font(run):
            if apply_custom_font:
                if formatting_options.get("font_family"):
                    run.font.name = formatting_options["font_family"]
                if formatting_options.get("font_size"):
                    run.font.size = Pt(int(formatting_options["font_size"]))
            else:
                if original_font_info.get('name'):
                    run.font.name = original_font_info['name']
                if original_font_info.get('size'):
                    run.font.size = original_font_info['size']
            if original_font_info.get('italic') is not None:
                run.italic = original_font_info['italic']

        # Split text by markers: **bold** (regular project headers) and
        # %%red%% (external/non-SYSTRA experience headers — red, not bold)
        pattern = re.compile(r'\*\*(.*?)\*\*|%%(.*?)%%')
        last_end = 0
        for m in pattern.finditer(text):
            if m.start() > last_end:
                part = text[last_end:m.start()]
                if part:
                    run = paragraph.add_run(part)
                    apply_common_font(run)
                    if original_font_info.get('color'):
                        run.font.color.rgb = original_font_info['color']
                    run.bold = False
                    run.font.bold = False

            if m.group(1) is not None:  # **bold**
                part = m.group(1)
                if part:
                    bold_run = paragraph.add_run(part)
                    apply_common_font(bold_run)
                    if original_font_info.get('color'):
                        bold_run.font.color.rgb = original_font_info['color']
                    bold_run.bold = True
                    bold_run.font.bold = True
            else:  # %%red%%
                part = m.group(2)
                if part:
                    red_run = paragraph.add_run(part)
                    apply_common_font(red_run)
                    red_run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
                    red_run.bold = False
                    red_run.font.bold = False

            last_end = m.end()

        if last_end < len(text):
            part = text[last_end:]
            if part:
                run = paragraph.add_run(part)
                apply_common_font(run)
                if original_font_info.get('color'):
                    run.font.color.rgb = original_font_info['color']
                run.bold = False
                run.font.bold = False


    def _format_education(self, education_list):
        """Format education entries"""
        if not education_list:
            return ""
        
        formatted = []
        for edu in education_list:
            if isinstance(edu, dict):
                cert_text = []
                if edu.get("Certificate/Diploma"):
                    cert_text.append(edu["Certificate/Diploma"])
                if edu.get("School/University"):
                    cert_text.append(edu["School/University"])
                if edu.get("year"):
                    cert_text.append(edu["year"])
                formatted.append(", ".join(cert_text))
            else:
                formatted.append(str(edu))
        
        return "\n".join([f"● {edu}" for edu in formatted])

    ########################################################
    def _format_skills(self, experience_list):
        """Format skills entries"""
        if not experience_list:
            return ""
        
        formatted = []
        for exp in experience_list:
            # Add role and company
            if exp.get("technical_expertise"):
                skills_text = exp['technical_expertise']
                individual_skills = [skill.strip() for skill in skills_text.split(',') if skill.strip()]
                formatted.extend(individual_skills)
        
        # Take first 8 skills and join them
        first_8_skills = formatted[:8]
        display_skills = ", ".join(first_8_skills) 
        return display_skills
    ########################################################

    def _format_certifications(self, certifications_list):
        """Format certifications entries"""
        if not certifications_list:
            return ""
        
        formatted = []
        for cert in certifications_list:
            if isinstance(cert, dict):
                cert_text = []
                if cert.get("programme"):
                    cert_text.append(cert["programme"])
                if cert.get("entity"):
                    cert_text.append(f" ({cert['entity']})")
                if cert.get("date"):
                    cert_text.append(f", {cert['date']}")
                formatted.append("".join(cert_text))
            else:
                formatted.append(str(cert))
        
        return "\n".join([f"● {cert}" for cert in formatted])
    
    def _format_memberships(self, memberships_list):
        """Format professional membership entries"""
        if not memberships_list:
            return ""
        
        formatted = []
        for membership in memberships_list:
            if isinstance(membership, dict):
                membership_text = []
                if membership.get("programme"):
                    membership_text.append(membership["programme"])
                if membership.get("date"):
                    membership_text.append(f", {membership['date']}")
                formatted.append("".join(membership_text))
            else:
                formatted.append(str(membership))
        return "\n".join([f"● {membership}" for membership in formatted])

    def _format_professional_experience(self, experience_list, external_list=None):
        """Format professional experience entries, followed by external (non-SYSTRA) experience"""
        formatted = []

        for exp in experience_list or []:
            exp_parts = []
            header_parts = []
            # Add role and company
            if exp.get("experience_header"):
                header_parts.append(exp['experience_header'])
            if exp.get("client"):
                header_parts.append(exp['client'])
            if exp.get("duration_period_year"):
                header_parts.append(exp['duration_period_year'])

            # Join header parts with commas
            if header_parts:
                header_line = ", ".join(header_parts)
                exp_parts.append(f"**{header_line}**") # Mark for bold formatting
            
            # Add description
            if exp.get("project_description"):
                exp_parts.append(exp['project_description'])
            if exp.get("expert_mission_description"):
                description = exp['expert_mission_description']
                exp_parts.append(description)
            
            # Combine header and description
            if exp_parts:
                formatted.append("\n".join(exp_parts))

        # NEW: append external (non-SYSTRA) experience after regular SYSTRA
        # projects, with the header marked for red (not bold) formatting.
        for ext in external_list or []:
            ext_parts = []
            if ext.get("header"):
                ext_parts.append(f"%%{ext['header']}%%")
            if ext.get("description"):
                ext_parts.append(ext["description"])
            if ext_parts:
                formatted.append("\n".join(ext_parts))

        return "\n\n".join([f"{exp}" for exp in formatted]) 
    
    def _format_professional_experience_summary(self, experience_list):
        """Format professional experience entries"""
        if not experience_list:
            return ""
        
        formatted = []
        for exp in experience_list:
            exp_parts = []
            # Add role and company
            if exp.get("employee_company"):
                exp_parts.append(exp['employee_company'])
            if exp.get("start_date"):
                full_years = re.findall(r'\b(?:19|20)\d{2}\b', exp.get("start_date") or "")
                start_year = None
                end_year = None
                if full_years:
                    if len(full_years) >= 1:
                        start_year = full_years[0]
                    if len(full_years) >= 2:
                        end_year = full_years[1]
                    if end_year is None:
                        end_year = "Present"
                    time_period = f"{start_year} - {end_year}"
                    exp_parts.append(time_period)
                else:
                    exp_parts.append(exp.get("start_date"))
            
            # Combine header and description
            if exp_parts:
                formatted.append(", ".join(exp_parts))
        
        return "\n".join([f"●   {exp}" for exp in formatted]) 

    def _update_header_footer(self, section, formatting_options):
        """Update header and footer with optional client/opportunity info from tkinter interface"""
        client = formatting_options.get("client_name", "")
        opportunity = formatting_options.get("opportunity_name", "")
        # Update footer

        def replace_in_paragraph(paragraph):
            # Recompose full text from runs to handle split placeholders
            full_text = "".join(run.text for run in paragraph.runs) if paragraph.runs else paragraph.text
            
            if "{{CLIENT}}" in full_text or "{{OPPORTUNITY}}" in full_text:
                new_text = full_text.replace("{{CLIENT}}", client).replace("{{OPPORTUNITY}}", opportunity)
                
                # Replace directly in runs like we do in _replace_placeholders
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                    for r in paragraph.runs[1:]:
                        r.text = ""
                else:
                    paragraph.clear()
                    paragraph.add_run(new_text)

        def process_footer(footer_obj):
            if not footer_obj:
                return
            # Plain paragraphs
            for p in footer_obj.paragraphs:
                replace_in_paragraph(p)
            # Paragraphs inside tables
            for tbl in footer_obj.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_in_paragraph(p)

        # Primary footer
        process_footer(section.footer)
        if hasattr(section, "first_page_footer") and section.first_page_footer:
            process_footer(section.first_page_footer)
        if hasattr(section, "even_page_footer") and section.even_page_footer:
            process_footer(section.even_page_footer)
